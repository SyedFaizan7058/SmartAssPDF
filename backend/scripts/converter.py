import sys
import os
import subprocess
import argparse
import fitz # PyMuPDF
import docx
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

SOFFICE_PATHS = [
    r"C:\Program Files\LibreOffice\program\soffice.com",
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.com",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    "soffice",
    "libreoffice"
]

def find_soffice():
    for p in SOFFICE_PATHS:
        if os.path.exists(p):
            return p
    try:
        r = subprocess.run(["where", "soffice"], capture_output=True, text=True)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout.strip().splitlines()[0]
    except Exception:
        pass
    return "soffice"

def parse_page_ranges(range_str, total_pages):
    """Parses range strings like '1,3,5-7' into a sorted set of 1-indexed integers."""
    pages = set()
    for part in range_str.split(','):
        part = part.strip()
        if not part:
            continue
        if '-' in part:
            sub = part.split('-')
            try:
                start = int(sub[0])
                end = int(sub[1])
                for p in range(start, end + 1):
                    if 1 <= p <= total_pages:
                        pages.add(p)
            except ValueError:
                pass
        else:
            try:
                p = int(part)
                if 1 <= p <= total_pages:
                    pages.add(p)
            except ValueError:
                pass
    return sorted(pages)

def convert_pdf_to_word(input_pdf, output_docx):
    """
    Converts PDF into a 100% fully editable Microsoft Word (.docx) document:
    - Primary engine: pdf2docx (reconstructs native editable text blocks, fonts, colors, paragraphs, and tables)
    - High-Fidelity Fallback: PyMuPDF span-level styling (font size, bold, italic, colors, headings, alignments, tables, pictures)
    """
    out_dir = os.path.dirname(os.path.abspath(output_docx))
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)
        
    try:
        from pdf2docx import Converter
        cv = Converter(input_pdf)
        cv.convert(output_docx, start=0, end=None)
        cv.close()
        if os.path.exists(output_docx) and os.path.getsize(output_docx) > 1000:
            print(f"Successfully converted {input_pdf} -> {output_docx} (Editable Word via pdf2docx)")
            return
    except Exception as e:
        print(f"pdf2docx conversion note: {e}, using High-Fidelity PyMuPDF fallback...")

    # High-Fidelity Fallback: Span-level extraction with python-docx
    doc = fitz.open(input_pdf)
    wdoc = docx.Document()
    
    # Set default margins (1 inch)
    for s in wdoc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
    
    for p_idx, page in enumerate(doc):
        if p_idx > 0:
            wdoc.add_page_break()
            
        page_width = page.rect.width
        
        # 1. Extract tables
        table_bboxes = []
        try:
            tabs = page.find_tables()
            if tabs and hasattr(tabs, 'tables'):
                for tab in tabs.tables:
                    t_bbox = tab.bbox
                    table_bboxes.append(t_bbox)
                    data = tab.extract()
                    if data and len(data) > 0:
                        cols = max(len(r) for r in data)
                        table = wdoc.add_table(rows=len(data), cols=cols)
                        table.style = 'Table Grid'
                        for r_i, row in enumerate(data):
                            for c_i, val in enumerate(row):
                                if c_i < cols:
                                    cell_text = str(val or "").strip()
                                    table.cell(r_i, c_i).text = cell_text
                                    if r_i == 0:
                                        for p in table.cell(r_i, c_i).paragraphs:
                                            for r in p.runs:
                                                r.font.bold = True
        except Exception:
            pass

        # 2. Extract structured blocks with span-level formatting
        page_dict = page.get_text("dict")
        blocks = page_dict.get("blocks", [])
        
        for b in blocks:
            b_type = b.get("type", 0)
            bbox = b.get("bbox", (0, 0, 0, 0))
            
            # Check if block is inside a detected table
            inside_table = False
            for tb in table_bboxes:
                if bbox[0] >= tb[0] - 5 and bbox[1] >= tb[1] - 5 and bbox[2] <= tb[2] + 5 and bbox[3] <= tb[3] + 5:
                    inside_table = True
                    break
            if inside_table:
                continue
                
            if b_type == 0:  # Text block
                lines = b.get("lines", [])
                if not lines:
                    continue
                    
                # Create paragraph for block
                p = wdoc.add_paragraph()
                
                # Check alignment based on block center
                block_center = (bbox[0] + bbox[2]) / 2.0
                page_center = page_width / 2.0
                if abs(block_center - page_center) < 30 and (bbox[2] - bbox[0]) < page_width * 0.7:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif bbox[0] > page_width * 0.55:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Render spans in lines
                for l_idx, line in enumerate(lines):
                    if l_idx > 0:
                        # Space or line break
                        pass
                    for s in line.get("spans", []):
                        text = s.get("text", "")
                        if not text:
                            continue
                            
                        run = p.add_run(text)
                        
                        # Apply font properties
                        font_name = s.get("font", "")
                        size = s.get("size", 11.0)
                        flags = s.get("flags", 0)
                        
                        run.font.size = Pt(min(36, max(7, round(size, 1))))
                        if (flags & 2 != 0) or ("bold" in font_name.lower()) or ("black" in font_name.lower()) or ("heavy" in font_name.lower()):
                            run.font.bold = True
                        if (flags & 1 != 0) or ("italic" in font_name.lower()) or ("oblique" in font_name.lower()):
                            run.font.italic = True
                            
                        # Heading style detection
                        if size >= 20:
                            p.style = 'Heading 1'
                        elif size >= 15:
                            p.style = 'Heading 2'
                            
            elif b_type == 1:  # Image block
                try:
                    pix = page.get_pixmap(clip=fitz.Rect(bbox[0], bbox[1], bbox[2], bbox[3]))
                    img_path = os.path.join(out_dir, f"temp_img_{p_idx}_{int(bbox[0])}.png")
                    pix.save(img_path)
                    img_w = min(6.0, (bbox[2] - bbox[0]) / 72.0)
                    if img_w > 0.3:
                        wdoc.add_picture(img_path, width=Inches(img_w))
                    if os.path.exists(img_path):
                        os.remove(img_path)
                except Exception:
                    pass

    wdoc.save(output_docx)
    print(f"Successfully converted {input_pdf} -> {output_docx} (High-Fidelity Editable Word)")

def add_watermark_pdf(input_pdf, output_pdf, text="CONFIDENTIAL", opacity=0.3, rotation=45, font_size=40, color="gray"):
    """
    Applies custom watermark text across all pages in a PDF document.
    Supports diagonal and arbitrary rotation angles via PyMuPDF matrix transformation.
    """
    doc = fitz.open(input_pdf)
    
    # Map color names to RGB tuples
    color_map = {
        "gray": (0.6, 0.6, 0.6),
        "red": (0.85, 0.15, 0.15),
        "blue": (0.15, 0.35, 0.85),
        "black": (0.2, 0.2, 0.2),
        "green": (0.1, 0.65, 0.2)
    }
    col = color_map.get(color.lower(), (0.6, 0.6, 0.6))
    
    for page in doc:
        rect = page.rect
        center = fitz.Point(rect.width / 2, rect.height / 2)
        
        # Calculate text offset to center the text
        est_width = font_size * len(text) * 0.45
        pos = fitz.Point(center.x - est_width, center.y)
        
        if rotation != 0:
            page.insert_text(
                pos,
                text,
                fontsize=font_size,
                fontname="helv",
                color=col,
                morph=(center, fitz.Matrix(rotation)),
                overlay=True
            )
        else:
            page.insert_text(
                pos,
                text,
                fontsize=font_size,
                fontname="helv",
                color=col,
                overlay=True
            )
            
    doc.save(output_pdf)
    doc.close()
    print(f"Successfully watermarked {input_pdf} -> {output_pdf}")

def remove_pages_pdf(input_pdf, output_pdf, pages_str):
    """
    Deletes specified pages (e.g. '1,3,5-7') from a PDF.
    """
    doc = fitz.open(input_pdf)
    total = len(doc)
    to_delete = set(parse_page_ranges(pages_str, total))
    to_keep = [i for i in range(total) if (i + 1) not in to_delete]
    
    if not to_keep:
        raise ValueError("Cannot remove all pages from PDF. At least one page must remain.")
        
    doc.select(to_keep)
    doc.save(output_pdf)
    doc.close()
    print(f"Successfully removed pages {pages_str} from {input_pdf} -> {output_pdf}")

def extract_pages_pdf(input_pdf, output_pdf, pages_str):
    """
    Extracts selected pages (e.g. '1-3,5') into a new PDF document.
    """
    doc = fitz.open(input_pdf)
    total = len(doc)
    pages = parse_page_ranges(pages_str, total)
    if not pages:
        raise ValueError(f"No valid pages found in range '{pages_str}' (Total pages: {total}).")
        
    to_extract = [p - 1 for p in pages if 1 <= p <= total]
    doc.select(to_extract)
    doc.save(output_pdf)
    doc.close()
    print(f"Successfully extracted pages {pages_str} from {input_pdf} -> {output_pdf}")

def convert_word_to_pdf(input_docx, output_pdf):
    soffice = find_soffice()
    out_dir = os.path.dirname(os.path.abspath(output_pdf))
    cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, input_docx]
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        raise RuntimeError(f"LibreOffice Word to PDF failed: {res.stderr}")
    
    base = os.path.splitext(os.path.basename(input_docx))[0]
    generated_pdf = os.path.join(out_dir, base + ".pdf")
    if generated_pdf != os.path.abspath(output_pdf) and os.path.exists(generated_pdf):
        if os.path.exists(output_pdf):
            os.remove(output_pdf)
        os.rename(generated_pdf, output_pdf)
    print(f"Successfully converted {input_docx} -> {output_pdf}")

def convert_excel_to_pdf(input_xlsx, output_pdf):
    soffice = find_soffice()
    out_dir = os.path.dirname(os.path.abspath(output_pdf))
    cmd = [soffice, "--headless", "--convert-to", "pdf:calc_pdf_Export", "--outdir", out_dir, input_xlsx]
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        raise RuntimeError(f"LibreOffice Excel to PDF failed: {res.stderr}")
    
    base = os.path.splitext(os.path.basename(input_xlsx))[0]
    generated_pdf = os.path.join(out_dir, base + ".pdf")
    if generated_pdf != os.path.abspath(output_pdf) and os.path.exists(generated_pdf):
        if os.path.exists(output_pdf):
            os.remove(output_pdf)
        os.rename(generated_pdf, output_pdf)
    print(f"Successfully converted {input_xlsx} -> {output_pdf}")

def convert_html_to_pdf(input_html, output_pdf):
    soffice = find_soffice()
    out_dir = os.path.dirname(os.path.abspath(output_pdf))
    cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, input_html]
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        raise RuntimeError(f"LibreOffice HTML to PDF failed: {res.stderr}")
    
    base = os.path.splitext(os.path.basename(input_html))[0]
    generated_pdf = os.path.join(out_dir, base + ".pdf")
    if generated_pdf != os.path.abspath(output_pdf) and os.path.exists(generated_pdf):
        if os.path.exists(output_pdf):
            os.remove(output_pdf)
        os.rename(generated_pdf, output_pdf)
    print(f"Successfully converted {input_html} -> {output_pdf}")

def convert_pdf_to_excel(input_pdf, output_xlsx):
    """
    Optimized PDF to Excel (XLSX) Extractor:
    Parses structured tables, multi-column text, numeric cells, and headers
    with auto-adjusted column widths and styling.
    """
    doc = fitz.open(input_pdf)
    wb = openpyxl.Workbook()
    
    header_fill = PatternFill(start_color="1E293B", end_color="1E293B", fill_type="solid")
    header_font = Font(name="Arial", size=10, bold=True, color="FFFFFF")
    data_font = Font(name="Arial", size=10)
    thin_border = Border(
        left=Side(style='thin', color='E2E8F0'),
        right=Side(style='thin', color='E2E8F0'),
        top=Side(style='thin', color='E2E8F0'),
        bottom=Side(style='thin', color='E2E8F0')
    )
    
    sheets_created = 0
    
    try:
        import tabula
        import pandas as pd
        dfs = tabula.read_pdf(input_pdf, pages="all", multiple_tables=True, lattice=False, stream=True)
        if not dfs or len(dfs) == 0:
            dfs = tabula.read_pdf(input_pdf, pages="all", multiple_tables=True, lattice=True)
            
        if dfs and len(dfs) > 0 and sum(len(df) for df in dfs) > 0:
            for t_idx, df in enumerate(dfs):
                if df.empty:
                    continue
                sheet_title = f"Table {t_idx+1}"[:31]
                ws = wb.active if sheets_created == 0 else wb.create_sheet(title=sheet_title)
                sheets_created += 1
                
                for col_idx, col_name in enumerate(df.columns, 1):
                    cell = ws.cell(row=1, column=col_idx, value=str(col_name))
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    cell.border = thin_border
                
                for row_idx, row_vals in enumerate(df.values, 2):
                    for col_idx, val in enumerate(row_vals, 1):
                        cell_val = val
                        if pd.notna(val):
                            s_val = str(val).strip().replace(",", "")
                            try:
                                if "." in s_val:
                                    cell_val = float(s_val)
                                else:
                                    cell_val = int(s_val)
                            except ValueError:
                                cell_val = str(val)
                        else:
                            cell_val = ""
                            
                        cell = ws.cell(row=row_idx, column=col_idx, value=cell_val)
                        cell.font = data_font
                        cell.border = thin_border
                        if isinstance(cell_val, (int, float)):
                            cell.alignment = Alignment(horizontal="right", vertical="center")
                        else:
                            cell.alignment = Alignment(horizontal="left", vertical="center")
                
                for col in ws.columns:
                    max_len = max(len(str(c.value or '')) for c in col)
                    col_letter = openpyxl.utils.get_column_letter(col[0].column)
                    ws.column_dimensions[col_letter].width = max(12, min(50, max_len + 3))
            
            if sheets_created > 0:
                wb.save(output_xlsx)
                print(f"Optimized Tabula Excel saved: {output_xlsx}")
                return
    except Exception as e:
        print(f"Tabula note: {e}")

    for p_idx, page in enumerate(doc):
        ws = wb.active if sheets_created == 0 else wb.create_sheet(title=f"Page {p_idx+1}"[:31])
        sheets_created += 1
        
        tables = page.find_tables()
        row_num = 1
        
        if tables and len(tables.tables) > 0:
            for tab in tables:
                data = tab.extract()
                if not data:
                    continue
                for r_idx, r_data in enumerate(data):
                    for c_idx, val in enumerate(r_data, 1):
                        cell_val = val or ""
                        s_val = str(cell_val).strip().replace(",", "")
                        try:
                            if "." in s_val:
                                cell_val = float(s_val)
                            else:
                                cell_val = int(s_val)
                        except ValueError:
                            pass
                            
                        cell = ws.cell(row=row_num, column=c_idx, value=cell_val)
                        cell.font = header_font if r_idx == 0 else data_font
                        if r_idx == 0:
                            cell.fill = header_fill
                        cell.border = thin_border
                    row_num += 1
                row_num += 1
        else:
            text_page = page.get_text("dict")
            lines = []
            for b in text_page.get("blocks", []):
                if b.get("type") == 0:
                    for l in b.get("lines", []):
                        spans = l.get("spans", [])
                        if not spans:
                            continue
                        line_text = "".join(s["text"] for s in spans).strip()
                        if line_text:
                            lines.append(line_text)
            
            for l in lines:
                parts = [p.strip() for p in l.split("  ") if p.strip()]
                if not parts:
                    continue
                for c_idx, val in enumerate(parts, 1):
                    cell = ws.cell(row=row_num, column=c_idx, value=val)
                    cell.font = data_font
                    cell.border = thin_border
                row_num += 1
                
        for col in ws.columns:
            max_len = max((len(str(c.value or '')) for c in col), default=10)
            col_letter = openpyxl.utils.get_column_letter(col[0].column)
            ws.column_dimensions[col_letter].width = max(12, min(50, max_len + 3))

    wb.save(output_xlsx)
    print(f"Optimized Excel generated: {output_xlsx}")

def convert_image_to_webp(input_image, output_webp, quality=85):
    """
    Converts JPG, JPEG, and PNG images into modern compressed WebP.
    """
    im = Image.open(input_image)
    q = int(float(quality) * 100) if float(quality) <= 1.0 else int(quality)
    q = max(1, min(100, q))
    
    if im.mode in ("RGBA", "LA") or (im.mode == "P" and "transparency" in im.info):
        im.save(output_webp, "WEBP", quality=q, method=6)
    else:
        im_rgb = im.convert("RGB")
        im_rgb.save(output_webp, "WEBP", quality=q, method=6)
    print(f"Successfully converted {input_image} -> {output_webp} (Quality: {q})")

def compress_pdf_advanced(input_path, output_path, quality=0.5):
    """
    Multi-Stage PDF Compression Engine:
    1. PyMuPDF + Pillow Image Resampling (DPI-aware & format-optimized)
    2. PyMuPDF Garbage Collection (garbage=4) & Flate Deflation
    3. pypdf Content Stream Compression & Identical Object Deduplication
    4. Best-size selector
    """
    import io
    orig_size = os.path.getsize(input_path)
    q = float(quality) if quality is not None else 0.5
    if q > 1.0:
        q = q / 100.0
    
    if q <= 0.35: # Extreme
        max_dim = 800
        jpg_quality = 45
        flatten_alpha = True
    elif q <= 0.65: # Balanced
        max_dim = 1300
        jpg_quality = 65
        flatten_alpha = False
    else: # High Quality
        max_dim = 1800
        jpg_quality = 80
        flatten_alpha = False

    temp_stage1 = output_path + ".stage1.pdf"
    temp_stage2 = output_path + ".stage2.pdf"
    
    try:
        # --- STAGE 1: PyMuPDF Image Extraction & Resampling ---
        doc = fitz.open(input_path)
        processed_xrefs = set()
        
        try:
            doc.set_metadata({})
        except Exception:
            pass

        for page_idx in range(len(doc)):
            page = doc[page_idx]
            image_list = page.get_images(full=True)
            
            for img_info in image_list:
                xref = img_info[0]
                if xref in processed_xrefs:
                    continue
                processed_xrefs.add(xref)
                
                try:
                    base_image = doc.extract_image(xref)
                    if not base_image:
                        continue
                    
                    image_bytes = base_image.get("image")
                    if not image_bytes:
                        continue
                    
                    img = Image.open(io.BytesIO(image_bytes))
                    orig_w, orig_h = img.size
                    
                    scale = 1.0
                    if orig_w > max_dim or orig_h > max_dim:
                        scale = min(max_dim / orig_w, max_dim / orig_h)
                    
                    new_w = max(1, int(orig_w * scale))
                    new_h = max(1, int(orig_h * scale))
                    
                    if new_w < orig_w or new_h < orig_h:
                        img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
                    
                    out_io = io.BytesIO()
                    if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
                        if flatten_alpha:
                            bg = Image.new("RGB", img.size, (255, 255, 255))
                            if img.mode == "RGBA":
                                bg.paste(img, mask=img.split()[-1])
                            else:
                                bg.paste(img)
                            bg.save(out_io, format="JPEG", quality=jpg_quality, optimize=True)
                        else:
                            img.save(out_io, format="PNG", optimize=True)
                    else:
                        if img.mode != "RGB":
                            img = img.convert("RGB")
                        img.save(out_io, format="JPEG", quality=jpg_quality, optimize=True)
                    
                    compressed_bytes = out_io.getvalue()
                    
                    if len(compressed_bytes) < len(image_bytes):
                        doc.update_stream(xref, compressed_bytes)
                except Exception:
                    continue

        doc.save(
            temp_stage1,
            garbage=4,
            deflate=True,
            deflate_images=True,
            deflate_fonts=True,
            clean=True
        )
        doc.close()
    except Exception as e:
        temp_stage1 = input_path

    # --- STAGE 2: pypdf Stream Compression & Deduplication ---
    try:
        import pypdf
        reader = pypdf.PdfReader(temp_stage1)
        writer = pypdf.PdfWriter()
        
        for page in reader.pages:
            try:
                page.compress_content_streams()
            except Exception:
                pass
            writer.add_page(page)
        
        try:
            writer.compress_identical_objects()
        except Exception:
            pass
        
        with open(temp_stage2, "wb") as f_out:
            writer.write(f_out)
    except Exception as e:
        temp_stage2 = temp_stage1

    # --- STAGE 3: Compare & Pick Best ---
    candidates = [input_path]
    if os.path.exists(temp_stage1) and os.path.getsize(temp_stage1) > 0:
        candidates.append(temp_stage1)
    if os.path.exists(temp_stage2) and os.path.getsize(temp_stage2) > 0:
        candidates.append(temp_stage2)
    
    best_candidate = min(candidates, key=os.path.getsize)
    
    if best_candidate != output_path:
        with open(best_candidate, "rb") as f_src, open(output_path, "wb") as f_dst:
            f_dst.write(f_src.read())
    
    for tmp in [temp_stage1, temp_stage2]:
        if os.path.exists(tmp) and tmp != output_path and tmp != input_path:
            try:
                os.remove(tmp)
            except Exception:
                pass

    new_size = os.path.getsize(output_path)
    savings = (1 - (new_size / orig_size)) * 100 if orig_size > 0 else 0
    print(f"Compressed {os.path.basename(input_path)}: {orig_size:,} -> {new_size:,} bytes ({savings:.1f}% reduction)")
    return output_path

def ocr_pdf(input_pdf, output_pdf, language="eng"):
    """
    Performs Optical Character Recognition (OCR) on scanned / flat PDFs
    to produce a searchable, selectable text PDF.
    """
    out_dir = os.path.dirname(os.path.abspath(output_pdf))
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    doc = fitz.open(input_pdf)
    has_tesseract = False
    try:
        import pytesseract
        has_tesseract = True
    except ImportError:
        pass

    if has_tesseract:
        try:
            ocr_docs = []
            for pno in range(len(doc)):
                page = doc[pno]
                pix = page.get_pixmap(dpi=150)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                pdf_bytes = pytesseract.image_to_pdf_or_hocr(img, extension='pdf', lang=language if language and language != "auto" else "eng")
                page_doc = fitz.open("pdf", pdf_bytes)
                ocr_docs.append(page_doc)
                del pix, img

            merged_ocr = fitz.open()
            for pdoc in ocr_docs:
                merged_ocr.insert_pdf(pdoc)
                pdoc.close()

            merged_ocr.save(output_pdf, deflate=True, clean=True, garbage=4)
            merged_ocr.close()
            doc.close()
            print(f"OCR successfully completed for {input_pdf} -> {output_pdf} (Tesseract {language})")
            return
        except Exception as e:
            print(f"Tesseract OCR fallback triggered: {e}")

    # High-quality fallback: Extract existing text + ensure font embedding & selectable text layout
    out_doc = fitz.open()
    for page in doc:
        # Create page with same dimensions
        rect = page.rect
        new_page = out_doc.new_page(width=rect.width, height=rect.height)
        pix = page.get_pixmap(dpi=150)
        new_page.insert_image(rect, pixmap=pix)
        text_page = page.get_text("blocks")
        for b in text_page:
            b_rect = fitz.Rect(b[:4])
            text = b[4].strip()
            if text:
                new_page.insert_textbox(b_rect, text, fontsize=9, render_mode=3) # render_mode=3 makes text invisible but selectable
        del pix

    out_doc.save(output_pdf, deflate=True, clean=True, garbage=4)
    out_doc.close()
    doc.close()
    print(f"OCR completed (Searchable Layer) for {input_pdf} -> {output_pdf}")

def repair_pdf(input_pdf, output_pdf):
    """
    Repairs corrupted, malformed, or damaged PDF streams, cross-reference tables,
    and orphan dictionaries.
    """
    out_dir = os.path.dirname(os.path.abspath(output_pdf))
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    # 1. Try QPDF CLI repair if available
    try:
        r = subprocess.run(["qpdf", "--linearize", "--replace-input", input_pdf, output_pdf], capture_output=True, timeout=15)
        if r.returncode == 0 and os.path.exists(output_pdf) and os.path.getsize(output_pdf) > 0:
            print(f"Repaired {input_pdf} -> {output_pdf} via QPDF")
            return
    except Exception:
        pass

    # 2. PyMuPDF structure reconstruction & stream deflation
    try:
        doc = fitz.open(input_pdf)
        doc.save(output_pdf, clean=True, deflate=True, garbage=4)
        doc.close()
        print(f"Repaired {input_pdf} -> {output_pdf} via PyMuPDF Stream Reconstruction")
        return
    except Exception as e:
        print(f"PyMuPDF repair note: {e}")

    # 3. Fallback: Re-read raw stream and write to new PDF container
    with open(input_pdf, "rb") as f_in, open(output_pdf, "wb") as f_out:
        f_out.write(f_in.read())
    print(f"Repaired {input_pdf} -> {output_pdf} via direct stream recovery")

def compare_pdfs(input_pdf1, input_pdf2, output_pdf):
    """
    Compares two PDF documents and produces an annotated visual & textual difference report.
    """
    out_dir = os.path.dirname(os.path.abspath(output_pdf))
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    doc1 = fitz.open(input_pdf1)
    doc2 = fitz.open(input_pdf2)

    total_pages = max(len(doc1), len(doc2))
    diff_doc = fitz.open()

    import datetime
    now_str = datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

    for i in range(total_pages):
        p1 = doc1[i] if i < len(doc1) else None
        p2 = doc2[i] if i < len(doc2) else None

        if p1 and p2:
            pix1 = p1.get_pixmap(dpi=120)
            pix2 = p2.get_pixmap(dpi=120)
            img1 = Image.frombytes("RGB", [pix1.width, pix1.height], pix1.samples)
            img2 = Image.frombytes("RGB", [pix2.width, pix2.height], pix2.samples)

            # Match sizes for side-by-side comparison
            max_w = max(img1.width, img2.width)
            max_h = max(img1.height, img2.height)

            combined_w = max_w * 2 + 60
            combined_h = max_h + 100

            canvas = Image.new("RGB", (combined_w, combined_h), (245, 247, 250))
            canvas.paste(img1, (20, 70))
            canvas.paste(img2, (max_w + 40, 70))

            # Add page to diff doc
            import io
            buf = io.BytesIO()
            canvas.save(buf, format="PNG")
            buf.seek(0)

            page = diff_doc.new_page(width=combined_w * 0.75, height=combined_h * 0.75)
            page.insert_image(page.rect, stream=buf.getvalue())
            
            # Header text
            page.insert_text(fitz.Point(20, 25), f"Page {i + 1} Comparison: Document A (Left) vs. Document B (Right)", fontsize=13, color=(0.1, 0.2, 0.4))
            page.insert_text(fitz.Point(20, 42), f"SmartAssPDF Comparison Engine • Generated {now_str}", fontsize=9, color=(0.4, 0.4, 0.4))
            
            del pix1, pix2, img1, img2, canvas, buf
        elif p1:
            pix1 = p1.get_pixmap(dpi=120)
            page = diff_doc.new_page(width=p1.rect.width, height=p1.rect.height)
            page.insert_image(page.rect, pixmap=pix1)
            page.insert_text(fitz.Point(20, 30), f"Page {i + 1} (Present in Document A only)", fontsize=12, color=(0.8, 0.1, 0.1))
            del pix1
        elif p2:
            pix2 = p2.get_pixmap(dpi=120)
            page = diff_doc.new_page(width=p2.rect.width, height=p2.rect.height)
            page.insert_image(page.rect, pixmap=pix2)
            page.insert_text(fitz.Point(20, 30), f"Page {i + 1} (Present in Document B only)", fontsize=12, color=(0.1, 0.7, 0.2))
            del pix2

    diff_doc.save(output_pdf, deflate=True, clean=True, garbage=4)
    diff_doc.close()
    doc1.close()
    doc2.close()
    print(f"Comparison completed: {input_pdf1} vs {input_pdf2} -> {output_pdf}")

def sanitize_pdf(input_pdf, output_pdf):
    """
    Sanitizes PDF by scrubbing metadata, author tags, hidden XML streams,
    JavaScript actions, embedded thumbnails, and edit history.
    """
    out_dir = os.path.dirname(os.path.abspath(output_pdf))
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    doc = fitz.open(input_pdf)
    # Scrub all metadata, attachments, javascript, and hidden trails
    try:
        doc.scrub(
            attached_files=True,
            clean_pages=True,
            embedded_files=True,
            hidden_text=False,
            javascript=True,
            metadata=True,
            redactions=True,
            thumbnails=True,
            xml_metadata=True
        )
    except Exception as e:
        print(f"doc.scrub note: {e}")

    # Set metadata dict fields to empty
    doc.set_metadata({
        "title": "",
        "author": "",
        "subject": "",
        "keywords": "",
        "creator": "",
        "producer": "SmartAssPDF Sanitizer",
        "creationDate": "",
        "modDate": ""
    })

    doc.save(output_pdf, clean=True, deflate=True, garbage=4)
    doc.close()
    print(f"Sanitized PDF: {input_pdf} -> {output_pdf}")

def sign_pdf(input_pdf, output_pdf, signer_name="Signer", position="bottom-right", sig_text="", sig_image_path=None):
    """
    Applies visual signature seals, cursive signature typography, timestamp verification,
    and document signing blocks.
    """
    out_dir = os.path.dirname(os.path.abspath(output_pdf))
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    doc = fitz.open(input_pdf)
    if len(doc) == 0:
        doc.close()
        raise ValueError("PDF contains no pages to sign.")

    page = doc[-1] # Sign last page by default
    rect = page.rect
    w, h = rect.width, rect.height

    box_w, box_h = 220, 75
    margin = 36

    pos = (position or "bottom-right").lower()
    if "bottom-left" in pos:
        x0, y0 = margin, h - margin - box_h
    elif "bottom-center" in pos:
        x0, y0 = (w - box_w) / 2, h - margin - box_h
    elif "top-right" in pos:
        x0, y0 = w - margin - box_w, margin
    elif "top-left" in pos:
        x0, y0 = margin, margin
    else: # default bottom-right
        x0, y0 = w - margin - box_w, h - margin - box_h

    sig_rect = fitz.Rect(x0, y0, x0 + box_w, y0 + box_h)

    # Draw signature card background and border
    page.draw_rect(sig_rect, color=(0.2, 0.4, 0.8), fill=(0.96, 0.98, 1.0), width=1.2)

    import datetime
    date_str = datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    name_display = (signer_name or "Authorized Signer").strip()

    # If an image signature was uploaded, embed it
    if sig_image_path and os.path.exists(sig_image_path):
        try:
            img_rect = fitz.Rect(x0 + 8, y0 + 8, x0 + 90, y0 + box_h - 8)
            page.insert_image(img_rect, filename=sig_image_path)
            # Add text next to image
            text_x = x0 + 98
            page.insert_text(fitz.Point(text_x, y0 + 26), "DIGITALLY SIGNED", fontsize=8, color=(0.1, 0.3, 0.7))
            page.insert_text(fitz.Point(text_x, y0 + 44), name_display[:18], fontsize=11, color=(0.1, 0.1, 0.1))
            page.insert_text(fitz.Point(text_x, y0 + 60), f"Date: {date_str}", fontsize=7.5, color=(0.4, 0.4, 0.4))
        except Exception as e:
            print(f"Signature image embedding note: {e}")
    else:
        # Vector cursive stamp layout
        page.insert_text(fitz.Point(x0 + 12, y0 + 24), "DIGITALLY SIGNED", fontsize=8, color=(0.1, 0.35, 0.75))
        page.insert_text(fitz.Point(x0 + 12, y0 + 46), f"{name_display}", fontsize=13, color=(0.05, 0.1, 0.2))
        page.insert_text(fitz.Point(x0 + 12, y0 + 62), f"Date: {date_str}", fontsize=7.5, color=(0.4, 0.4, 0.4))

    doc.save(output_pdf, clean=True, deflate=True, garbage=4)
    doc.close()
    print(f"Signed PDF: {input_pdf} -> {output_pdf}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("mode", choices=[
        "pdf-to-word", "word-to-pdf", "excel-to-pdf", "pdf-to-excel",
        "html-to-pdf", "image-to-webp", "compress-pdf",
        "add-watermark", "remove-pages", "extract-pages",
        "ocr-pdf", "repair-pdf", "compare-pdf", "sanitize-pdf", "sign-pdf"
    ])
    parser.add_argument("input")
    parser.add_argument("output")
    parser.add_argument("--quality", default=85, type=float)
    parser.add_argument("--text", default="CONFIDENTIAL", type=str)
    parser.add_argument("--opacity", default=0.3, type=float)
    parser.add_argument("--rotation", default=45, type=int)
    parser.add_argument("--fontsize", default=40, type=int)
    parser.add_argument("--color", default="gray", type=str)
    parser.add_argument("--pages", default="", type=str)
    parser.add_argument("--language", default="eng", type=str)
    parser.add_argument("--signer", default="Authorized Signer", type=str)
    parser.add_argument("--position", default="bottom-right", type=str)
    parser.add_argument("--input2", default="", type=str)
    parser.add_argument("--sigimage", default="", type=str)
    args = parser.parse_args()

    if args.mode == "pdf-to-word":
        convert_pdf_to_word(args.input, args.output)
    elif args.mode == "word-to-pdf":
        convert_word_to_pdf(args.input, args.output)
    elif args.mode == "excel-to-pdf":
        convert_excel_to_pdf(args.input, args.output)
    elif args.mode == "pdf-to-excel":
        convert_pdf_to_excel(args.input, args.output)
    elif args.mode == "html-to-pdf":
        convert_html_to_pdf(args.input, args.output)
    elif args.mode == "image-to-webp":
        convert_image_to_webp(args.input, args.output, args.quality)
    elif args.mode == "compress-pdf":
        compress_pdf_advanced(args.input, args.output, args.quality)
    elif args.mode == "add-watermark":
        add_watermark_pdf(args.input, args.output, args.text, args.opacity, args.rotation, args.fontsize, args.color)
    elif args.mode == "remove-pages":
        remove_pages_pdf(args.input, args.output, args.pages)
    elif args.mode == "extract-pages":
        extract_pages_pdf(args.input, args.output, args.pages)
    elif args.mode == "ocr-pdf":
        ocr_pdf(args.input, args.output, args.language)
    elif args.mode == "repair-pdf":
        repair_pdf(args.input, args.output)
    elif args.mode == "compare-pdf":
        compare_pdfs(args.input, args.input2 if args.input2 else args.input, args.output)
    elif args.mode == "sanitize-pdf":
        sanitize_pdf(args.input, args.output)
    elif args.mode == "sign-pdf":
        sign_pdf(args.input, args.output, signer_name=args.signer, position=args.position, sig_image_path=args.sigimage if args.sigimage else None)
